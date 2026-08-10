#!/usr/bin/env python3

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_font(style, name, size, bold=None):
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def main():
    if len(sys.argv) != 2:
        raise SystemExit("usage: create-tmpa-reference-docx.py OUTPUT.docx")
    output = Path(sys.argv[1])
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    set_font(styles["Normal"], "Noto Sans CJK SC", 10.5)
    set_font(styles["Title"], "Noto Sans CJK SC", 22, True)
    set_font(styles["Subtitle"], "Noto Sans CJK SC", 13)
    set_font(styles["Heading 1"], "Noto Sans CJK SC", 17, True)
    set_font(styles["Heading 2"], "Noto Sans CJK SC", 14, True)
    set_font(styles["Heading 3"], "Noto Sans CJK SC", 12, True)
    set_font(styles["Caption"], "Noto Sans CJK SC", 9)
    if "Footnote Text" in styles:
        set_font(styles["Footnote Text"], "Noto Sans CJK SC", 8.5)

    normal = styles["Normal"].paragraph_format
    normal.space_after = Pt(6)
    normal.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("TMPA Publication Dossier RC1  |  ")
    run.font.name = "Noto Sans CJK SC"
    run.font.size = Pt(8)
    page_field(footer)

    document.add_paragraph("TMPA reference layout", style="Title")
    document.add_paragraph("Generated for the bilingual RC1 publication dossier.")
    document.save(output)


if __name__ == "__main__":
    main()
